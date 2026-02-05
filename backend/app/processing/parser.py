import os
import io
import json
import re
import fitz  # PyMuPDF
from PIL import Image

def _assets_dir(file_path):
    base = os.path.splitext(os.path.basename(file_path))[0]
    assets_root = os.path.join(os.path.dirname(file_path), "extracted_assets", base)
    os.makedirs(assets_root, exist_ok=True)
    return assets_root

def _save_image_bytes(image_bytes, out_path):
    with open(out_path, "wb") as f:
        f.write(image_bytes)
    return out_path

def _ocr_image_bytes(image_bytes):
    try:
        import pytesseract
    except Exception as exc:
        raise RuntimeError("pytesseract is required for OCR processing") from exc

    img = Image.open(io.BytesIO(image_bytes))
    return pytesseract.image_to_string(img)

def _extract_chart_numbers(text):
    # Best-effort numeric extraction from OCR text
    return re.findall(r"[-+]?(?:\d+\.?\d*|\d*\.\d+)", text)

def _extract_pdf(pdf_path):
    import re
    try:
        import pdfplumber
    except Exception as exc:
        raise RuntimeError("pdfplumber is required for PDF table extraction") from exc

    assets_dir = _assets_dir(pdf_path)
    pages = []
    with fitz.open(pdf_path) as doc:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(doc):
                text_parts = [page.get_text()]

                # Table extraction via pdfplumber (best-effort)
                try:
                    pdf_page = pdf.pages[i]
                    tables = pdf_page.extract_tables() or []
                    for t_index, table in enumerate(tables, start=1):
                        clean_rows = [[(cell or "").strip() for cell in row] for row in table]
                        csv_lines = [", ".join(row) for row in clean_rows if any(row)]
                        table_json = json.dumps(clean_rows, ensure_ascii=True)
                        if csv_lines:
                            text_parts.append(f"[TABLE {t_index} CSV]\n" + "\n".join(csv_lines))
                        text_parts.append(f"[TABLE {t_index} JSON]\n{table_json}")
                except Exception:
                    pass

                # Image extraction + OCR + numeric candidates
                try:
                    for img_index, img in enumerate(page.get_images(full=True), start=1):
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n > 4:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_path = os.path.join(
                            assets_dir,
                            f"page_{i+1:03d}_img_{img_index:02d}.png",
                        )
                        pix.save(img_path)
                        with open(img_path, "rb") as f:
                            image_bytes = f.read()
                        ocr_text = _ocr_image_bytes(image_bytes).strip()
                        if ocr_text:
                            text_parts.append(f"[IMAGE OCR] {ocr_text}")
                            nums = _extract_chart_numbers(ocr_text)
                            if nums:
                                text_parts.append("[CHART DATA CANDIDATES] " + ", ".join(nums))
                except Exception:
                    pass

                pages.append({"page": i + 1, "text": "\n".join([p for p in text_parts if p])})
    return pages

def _ocr_pdf(pdf_path):
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except Exception as exc:
        raise RuntimeError("pdf2image and pytesseract are required for OCR processing") from exc

    images = convert_from_path(pdf_path, dpi=300)
    pages = []
    for i, img in enumerate(images):
        text = pytesseract.image_to_string(img)
        pages.append({"page": i + 1, "text": text})
    return pages

def _extract_docx(docx_path):
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX processing") from exc

    doc = Document(docx_path)
    parts = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    for t_index, table in enumerate(doc.tables, start=1):
        table_rows = []
        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                cell_parts = [cp.text.strip() for cp in cell.paragraphs if cp.text.strip()]
                cell_text = " ".join(cell_parts) if cell_parts else cell.text.strip()
                cell_texts.append(cell_text)
            table_rows.append(cell_texts)
            row_text = " | ".join([t for t in cell_texts if t])
            if row_text:
                parts.append(row_text)
        table_json = json.dumps(table_rows, ensure_ascii=True)
        parts.append(f"[TABLE {t_index} JSON]\n{table_json}")
        csv_lines = [", ".join(row) for row in table_rows if any(row)]
        if csv_lines:
            parts.append(f"[TABLE {t_index} CSV]\n" + "\n".join(csv_lines))

    # Inline images (charts/diagrams) -> save + OCR
    assets_dir = _assets_dir(docx_path)
    try:
        for i, shape in enumerate(doc.inline_shapes, start=1):
            try:
                rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                image_part = doc.part.related_parts[rId]
                img_path = os.path.join(assets_dir, f"docx_img_{i:02d}.png")
                _save_image_bytes(image_part.blob, img_path)
                ocr_text = _ocr_image_bytes(image_part.blob).strip()
                if ocr_text:
                    parts.append(f"[IMAGE OCR] {ocr_text}")
                    nums = _extract_chart_numbers(ocr_text)
                    if nums:
                        parts.append("[CHART DATA CANDIDATES] " + ", ".join(nums))
            except Exception:
                continue
    except Exception:
        pass

    text = "\n".join(parts)
    return [{"page": 1, "text": text}]

def _extract_doc(doc_path):
    try:
        import textract
    except Exception as exc:
        raise RuntimeError("textract is required for DOC processing") from exc

    text = textract.process(doc_path).decode("utf-8", errors="ignore")
    return [{"page": 1, "text": text}]

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        pages = _extract_pdf(file_path)
        total_text = "".join([p.get("text", "") for p in pages]).strip()
        if len(total_text) < 50:
            return _ocr_pdf(file_path)
        return pages
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext == ".doc":
        return _extract_doc(file_path)
    return []
